from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "Relatorio_Sprint2_RabbitMQ.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GRAY = "666666"
LIGHT_GRAY = "F2F4F7"
NOTE_BLUE = "E8EEF5"
BLACK = "000000"
WHITE = "FFFFFF"
CONTENT_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_dxa):
            set_cell_width(row.cells[index], width)
            set_cell_margins(row.cells[index])


def set_font(run, name="Calibri", size=11, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri", size=11, bold=None, color=BLACK) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=11, color=BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        set_style_font(style, size=size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, size=11, color=BLACK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, name="Consolas", size=8.5, color=BLACK)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.left_indent = Inches(0.12)
    code.paragraph_format.right_indent = Inches(0.12)

    note = styles.add_style("Note", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(note, size=10.5, color=BLACK)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.left_indent = Inches(0.15)
    note.paragraph_format.right_indent = Inches(0.15)


def add_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_name
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("FastDelivery | Sprint 2 | Integração MOM")
    set_font(run, size=8.5, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("FastDelivery - Relatório Sprint 2  |  Página ")
    set_font(run, size=8.5, color=GRAY)
    add_field(footer, "PAGE")


def add_para(doc: Document, text="", style=None, align=None, before=0, after=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        p.add_run(text)
    return p


def add_rich_para(doc: Document, parts, style=None, after=None):
    p = doc.add_paragraph(style=style)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    for text, kwargs in parts:
        run = p.add_run(text)
        if kwargs:
            set_font(run, **kwargs)
    return p


def add_title_block(doc: Document) -> None:
    p = add_para(doc, after=2)
    r = p.add_run("RELATÓRIO DE INTEGRAÇÃO MOM")
    set_font(r, size=22, bold=True, color=BLACK)
    p = add_para(doc, after=10)
    r = p.add_run("Sprint 2 - FastDelivery com RabbitMQ")
    set_font(r, size=14, bold=True, color=BLUE)

    rows = [
        ("Disciplina", "Laboratório de Desenvolvimento de Aplicações Móveis e Distribuídas"),
        ("Projeto", "FastDelivery"),
        ("Referência", "Seção 3.3, página 5 do enunciado Projeto_LAMD_60445_2026_1"),
        ("Escopo", "Integração assíncrona orientada a eventos com RabbitMQ"),
    ]
    for label, value in rows:
        p = add_para(doc, after=1)
        r = p.add_run(f"{label}: ")
        set_font(r, size=9.5, bold=True, color=GRAY)
        r = p.add_run(value)
        set_font(r, size=9.5, color=GRAY)


def add_heading(doc: Document, text: str, level=1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=font_size, bold=True, color=DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line or " ")
        set_font(run, name="Consolas", size=8.5, color=BLACK)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, NOTE_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Nota: ")
    set_font(r, size=9.5, bold=True, color=DARK_BLUE)
    r = p.add_run(text)
    set_font(r, size=9.5, color=BLACK)
    set_table_geometry(table, [CONTENT_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def first_page(doc: Document) -> None:
    add_title_block(doc)
    add_heading(doc, "Objetivo", 1)
    add_para(doc, "A Sprint 2 integrou um Middleware Orientado a Mensagens (MOM) ao backend Flask do FastDelivery. O objetivo foi permitir que fatos relevantes do domínio de entregas fossem publicados e processados de forma assíncrona, sem chamada REST direta entre produtor e consumidor.")

    add_heading(doc, "Decisões de design", 1)
    add_para(doc, "Foi adotado o RabbitMQ por ser um broker dedicado, com suporte nativo a filas duráveis, mensagens persistentes, confirmação de publicação, confirmação manual de consumo e Dead-Letter Queue (DLQ). O padrão aplicado é Publish-Subscribe sobre a exchange topic fastdelivery.events.")
    add_para(doc, "O backend Flask atua como produtor. Por padrão, o consumidor roda em outro processo com py consumer_worker.py. Assim, o produtor publica mesmo quando o consumidor está desligado; a fila fastdelivery.entregas armazena o backlog; quando o worker inicia, ele consome as mensagens pendentes.")

    add_heading(doc, "Fluxo implementado", 1)
    add_numbers(doc, [
        "POST /entregas cria uma entrega e publica entrega.criada.",
        "PATCH /entregas/<id>/status altera o estado e publica entrega.status_atualizado.",
    ])
    add_para(doc, "Os eventos seguem pela exchange fastdelivery.events para a fila durável fastdelivery.entregas. O worker processa a mensagem e somente então envia ack. Falhas recebem nack(requeue=False) e seguem para fastdelivery.dlq.")

    add_heading(doc, "Desafios e conclusão", 1)
    add_para(doc, "A primeira versão utilizava Redis. A revisão técnica motivou a migração para RabbitMQ e a inclusão de mensagens persistentes (delivery_mode=2), publisher confirms, uma tentativa de reconexão, DLQ e histórico idempotente no SQLite por evento_id. A solução atende ao enunciado: o broker é o ponto central, existem dois eventos de negócio documentados e o fluxo assíncrono pode ser demonstrado por logs, UI do RabbitMQ e GET /eventos.")


def requirement_page(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 1 - Conferência do enunciado", 1)
    add_para(doc, "A seção 3.3 do enunciado solicita MOM operacional com evidência, produtor e consumidor, publicação em dois momentos de negócio, catálogo dos eventos, prova de assincronicidade e um relatório de integração de uma página. A página anterior é o relatório principal; os anexos apoiam a demonstração.")
    add_table(doc,
              ["Requisito", "Implementação", "Evidência recomendada"],
              [
                  ["MOM operacional", "RabbitMQ via Docker Compose, AMQP 5672 e UI 15672", "UI do RabbitMQ e docker compose ps"],
                  ["Produtor e consumidor", "EntregaEventProducer e consumer_worker.py", "Terminais separados"],
                  ["Dois momentos", "Criação e alteração de status", "Logs dos dois tópicos"],
                  ["Catálogo", "Eventos, routing keys e payloads", "Anexo 2"],
                  ["Assincronicidade", "Backlog com worker desligado", "Fila Ready > 0 e depois Ready = 0"],
                  ["Relatório", "Síntese de decisões e desafios", "Primeira página"],
              ],
              [1800, 3900, 3660])

    add_heading(doc, "O que foi adicionado", 1)
    add_table(doc,
              ["Componente", "Responsabilidade"],
              [
                  ["app/mom/barramento.py", "Interface EventBus, RabbitMQEventBus e InMemoryEventBus para testes"],
                  ["app/mom/eventos.py", "Tópicos, payload comum e evento_id"],
                  ["app/mom/producer.py", "Publicação dos fatos do domínio"],
                  ["app/mom/consumer.py", "Handlers, histórico persistente e deduplicação"],
                  ["consumer_worker.py", "Consumidor standalone sem Flask ou REST"],
                  ["evento_processado_repository.py", "Histórico compartilhado no SQLite"],
                  ["docker-compose.yml", "RabbitMQ com UI e volume persistente"],
                  ["GET /eventos", "Consulta aos eventos processados"],
              ],
              [3150, 6210])

    add_heading(doc, "Garantias técnicas", 1)
    add_bullets(doc, [
        "Exchange topic fastdelivery.events e fila durável fastdelivery.entregas ligada com routing key #.",
        "Mensagens persistentes com delivery_mode=2, publisher confirms e mandatory=True.",
        "Ack manual após sucesso; nack(requeue=False) envia falhas para fastdelivery.dlq.",
        "Uma nova tentativa de publicação após falha de conexão.",
        "Deduplicação por evento_id no histórico SQLite.",
    ])


def catalog_page(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 2 - Catálogo de eventos", 1)
    add_heading(doc, "entrega.criada", 2)
    add_table(doc, ["Campo", "Valor"], [
        ["Momento", "Após EntregaUseCases.criar_entrega() persistir a nova entrega"],
        ["Produtor", "EntregaEventProducer.entrega_criada()"],
        ["Consumidor", "_ao_criar_entrega() executado por consumer_worker.py"],
        ["Exchange / routing key", "fastdelivery.events / entrega.criada"],
        ["Fila", "fastdelivery.entregas"],
    ], [2400, 6960])
    add_code(doc, """{
  "evento_id": "uuid-gerado-na-publicacao",
  "evento": "entrega.criada",
  "dados": {
    "id": 1, "descricao": "Pacote Sprint 2",
    "origem": "Rua A, 100", "destino": "Rua B, 200",
    "status": "pendente", "cliente_id": "cliente-001"
  },
  "timestamp": "2026-06-01T10:00:00"
}""")

    add_heading(doc, "entrega.status_atualizado", 2)
    add_table(doc, ["Campo", "Valor"], [
        ["Momento", "Após EntregaUseCases.atualizar_status() persistir o novo estado"],
        ["Produtor", "EntregaEventProducer.status_alterado()"],
        ["Consumidor", "_ao_alterar_status() executado por consumer_worker.py"],
        ["Exchange / routing key", "fastdelivery.events / entrega.status_atualizado"],
        ["Fila", "fastdelivery.entregas"],
    ], [2400, 6960])
    add_code(doc, """{
  "evento_id": "uuid-gerado-na-publicacao",
  "evento": "entrega.status_atualizado",
  "dados": {
    "id": 1, "status_anterior": "pendente",
    "status_novo": "aceito", "cliente_id": "cliente-001"
  },
  "timestamp": "2026-06-01T10:02:00"
}""")


def tests_page(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 3 - Testes executados", 1)
    add_para(doc, "Comando executado em Code/server/:")
    add_code(doc, "py -m unittest discover -s tests -v")
    add_para(doc, "Resultado obtido em 01/06/2026:")
    add_code(doc, "Ran 8 tests in 0.087s\nOK")
    add_table(doc, ["Teste", "Validação"], [
        ["Modo em memória explícito", "EVENT_BUS=in_memory cria o barramento reservado a testes"],
        ["Configuração padrão", "Sem variável específica, cria RabbitMQEventBus"],
        ["Ack de sucesso", "Handler bem-sucedido produz basic_ack"],
        ["Falha de handler", "Exceção produz basic_nack(requeue=False) para DLQ"],
        ["Tópico desconhecido", "Mensagem sem handler também segue para DLQ"],
        ["Canal encerrado", "A conexão de publicação é substituída"],
        ["Falha transitória", "O produtor repete a publicação uma vez"],
        ["Idempotência", "O mesmo evento_id é persistido uma única vez"],
    ], [3300, 6060])
    add_note(doc, "tests/test_entregas.py ainda é um placeholder. Os endpoints REST devem ser comprovados manualmente no roteiro a seguir. Não apresente a suíte MOM como se também cobrisse endpoints.")


def evidence_page_one(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 4 - Roteiro de execução e prints", 1)
    add_heading(doc, "1. Preparar o broker", 2)
    add_para(doc, "Abra o Docker Desktop. Depois execute em um PowerShell:")
    add_code(doc, r"""cd C:\Users\gpercope\Documents\GitHub\Lab-DAMD\Code\server
docker compose up -d
docker compose ps
Copy-Item .env.example .env -ErrorAction SilentlyContinue
pip install -r requirements.txt""")
    add_para(doc, "Abra http://localhost:15672 e entre com guest / guest.")
    add_note(doc, "Print E1 - Capture docker compose ps e, na UI, Queues and Streams com fastdelivery.entregas e fastdelivery.dlq.")

    add_heading(doc, "2. Subir apenas o backend produtor", 2)
    add_code(doc, r"""cd C:\Users\gpercope\Documents\GitHub\Lab-DAMD\Code\server
py main.py""")
    add_note(doc, "Print E2 - Capture o terminal do Flask mostrando Mecanismo MOM.: rabbitmq e Consumidor in-process: OFF.")

    add_heading(doc, "3. Publicar com o worker desligado", 2)
    add_code(doc, r"""$body = @{
  descricao = "Pacote demonstracao Sprint 2"
  origem = "Rua A, 100"
  destino = "Rua B, 200"
  cliente_id = "cliente-001"
} | ConvertTo-Json

$entrega = Invoke-RestMethod -Method Post `
  -Uri "http://localhost:5000/entregas" `
  -ContentType "application/json" -Body $body

Invoke-RestMethod -Method Patch `
  -Uri "http://localhost:5000/entregas/$($entrega.id)/status" `
  -ContentType "application/json" -Body '{"status":"aceito"}'""")
    add_note(doc, "Print E3 - Com o worker desligado, abra Queues and Streams > fastdelivery.entregas e capture Ready aumentando em 2. Este é o print principal da assincronicidade.")


def evidence_page_two(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 4 - Roteiro de execução e prints (continuação)", 1)
    add_heading(doc, "4. Ligar o consumidor independente", 2)
    add_code(doc, r"""cd C:\Users\gpercope\Documents\GitHub\Lab-DAMD\Code\server
py consumer_worker.py""")
    add_note(doc, "Print E4 - Capture os logs do worker para entrega.criada e entrega.status_atualizado. Print E5 - Atualize a UI e capture Ready = 0.")

    add_heading(doc, "5. Consultar o histórico", 2)
    add_code(doc, r"""Invoke-RestMethod -Uri "http://localhost:5000/eventos" |
  ConvertTo-Json -Depth 10""")
    add_note(doc, "Print E6 - Capture o JSON contendo os dois eventos, seus evento_id, tipos, mensagens e payloads.")

    add_heading(doc, "6. Executar a suíte automatizada", 2)
    add_code(doc, "py -m unittest discover -s tests -v")
    add_note(doc, "Print E7 - Capture a saída final com Ran 8 tests e OK.")

    add_heading(doc, "Evidências extras", 2)
    add_para(doc, "Persistência: com o worker desligado, publique uma entrega, confira Ready > 0 e execute:")
    add_code(doc, "docker compose restart rabbitmq\ndocker compose ps")
    add_para(doc, "Após o container ficar saudável, a mensagem deve continuar na fila. Capture antes e depois.")
    add_para(doc, "DLQ: com o worker ligado, publique um tópico sem handler:")
    add_code(doc, """py -c "from app.mom.barramento import criar_barramento; from app.mom.eventos import obter_payload; bus=criar_barramento(); bus.publicar('entrega.desconhecida', obter_payload('entrega.desconhecida', {'teste': True})); bus.parar()" """)
    add_para(doc, "Capture o log do worker e a fila fastdelivery.dlq com mensagem em Ready.")


def presentation_page(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "Anexo 5 - Roteiro curto para apresentação", 1)
    add_numbers(doc, [
        "Mostrar a UI do RabbitMQ e as filas fastdelivery.entregas e fastdelivery.dlq.",
        "Explicar que o Flask é produtor puro e que consumer_worker.py roda separado, sem Flask e sem REST.",
        "Manter o worker desligado, executar POST e PATCH e mostrar o backlog crescendo.",
        "Ligar o worker e mostrar os dois eventos sendo consumidos.",
        "Executar GET /eventos e mostrar o histórico persistente.",
        "Mostrar rapidamente a saída dos oito testes automatizados.",
        "Como evidência extra, demonstrar a DLQ ou a persistência após reinício do RabbitMQ.",
    ])
    add_heading(doc, "Aviso sobre evidências antigas", 1)
    add_note(doc, "Os arquivos em docs/Sprint2/evidencias foram gerados antes da migração e ainda descrevem Redis. Não use esses registros como prova atual da solução RabbitMQ. Produza capturas novas com o roteiro deste documento.")


def build() -> None:
    doc = Document()
    configure_styles(doc)
    setup_document(doc)
    first_page(doc)
    requirement_page(doc)
    catalog_page(doc)
    tests_page(doc)
    evidence_page_one(doc)
    evidence_page_two(doc)
    presentation_page(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
